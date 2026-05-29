# ================================
# SCRIPT DE PRE-PROCESAMIENTO DE DOCUMENTOS
# ================================
# Propósito: Convertir cualquier PDF/texto en chunks + embeddings + JSON
# Uso: python core/preprocess.py --input "ruta/al/documento.pdf" --output "ruta/output.json"
# ================================

import os
import sys
import json
import argparse
from datetime import datetime
from typing import List, Dict, Any
import hashlib

# ================================
# IMPORTS NUEVOS PARA MODO CONFIG
# ================================
import sys
import os

# Agregar raíz del proyecto al path para imports
project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
if project_root not in sys.path:
    sys.path.insert(0, project_root)

from utils.config_loader import load_book_config, get_config_value

# ================================
# DEPENDENCIAS (instalar con: pip install pypdf sentence-transformers)
# ================================
from pypdf import PdfReader
from sentence_transformers import SentenceTransformer

# ================================
# DEPENDENCIAS (instalar con: pip install -r requirements.txt)
# ================================
# requirements.txt:
# pypdf>=3.0.0
# sentence-transformers>=2.2.0
# einops>=0.6.0  ← NUEVA: requerida por nomic-embed-text

# ================================
# CONFIGURACIÓN
# ================================
CHUNK_SIZE = 512          # Tokens por chunk (ajustable: 256-1024)
CHUNK_OVERLAP = 50        # Tokens de overlap entre chunks
EMBEDDING_MODEL = "nomic-ai/nomic-embed-text-v1.5"  # Modelo de embeddings
MAX_CHUNKS_PER_SAVE = 1000  # Para no saturar memoria en documentos muy grandes

# ================================
# FUNCIONES AUXILIARES
# ================================

# ================================
# EXTRACCIÓN DE TEXTO (MULTI-FORMATO)
# ================================

def extract_text_from_file(file_path: str) -> str:
    """
    Extrae texto de un archivo PDF o DOCX automáticamente.
    Detecta el formato por la extensión del archivo.
    """
    print(f"📄 Extrayendo texto de: {file_path}")
    
    # Detectar formato por extensión
    file_ext = os.path.splitext(file_path)[1].lower()
    
    if file_ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif file_ext in ['.docx', '.doc']:
        return extract_text_from_docx(file_path)
    elif file_ext == '.txt':
        return extract_text_from_txt(file_path)
    else:
        raise ValueError(f"Formato no soportado: {file_ext}. Usa PDF, DOCX o TXT.")


def extract_text_from_pdf(pdf_path: str) -> str:
    """
    Extrae texto de un archivo PDF.
    """
    from pypdf import PdfReader
    
    reader = PdfReader(pdf_path)
    text = ""
    
    for i, page in enumerate(reader.pages, 1):
        page_text = page.extract_text()
        if page_text:
            text += f"\n[Página {i}]\n{page_text}"
    
    print(f"✅ Texto extraído: {len(text)} caracteres, {len(reader.pages)} páginas")
    return text


def extract_text_from_docx(docx_path: str) -> str:
    """
    Extrae texto de un archivo Word (.docx).
    """
    from docx import Document
    
    print(f"📄 Extrayendo texto de: {docx_path}")
    
    doc = Document(docx_path)
    text = ""
    
    for i, para in enumerate(doc.paragraphs, 1):
        if para.text.strip():
            text += f"\n[Párrafo {i}]\n{para.text}"
    
    # También extraer texto de tablas si existen
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += f"\n{cell.text}"
    
    print(f"✅ Texto extraído: {len(text)} caracteres, {len(doc.paragraphs)} párrafos")
    return text


def extract_text_from_txt(txt_path: str) -> str:
    """
    Extrae texto de un archivo TXT plano.
    """
    print(f"📄 Extrayendo texto de: {txt_path}")
    
    with open(txt_path, 'r', encoding='utf-8') as f:
        text = f.read()
    
    print(f"✅ Texto extraído: {len(text)} caracteres")
    return text

def clean_text(text: str) -> str:
    """
    Limpieza básica del texto:
    - Eliminar saltos de línea excesivos
    - Unificar espacios
    - Eliminar caracteres no imprimibles
    """
    print("🧹 Limpiando texto...")
    
    # Eliminar múltiples saltos de línea
    import re
    text = re.sub(r'\n\s*\n', '\n\n', text)
    
    # Unificar espacios
    text = re.sub(r' +', ' ', text)
    
    # Eliminar caracteres no imprimibles (excepto saltos normales)
    text = ''.join(char for char in text if char.isprintable() or char in '\n')
    
    print(f"✅ Texto limpio: {len(text)} caracteres")
    return text


def create_chunks(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[Dict[str, Any]]:
    """
    Divide el texto en chunks por tamaño (no solo por párrafos).
    Cada chunk incluye metadata para rastreo.
    """
    print(f"✂️ Creando chunks (tamaño: {chunk_size} caracteres, overlap: {overlap})...")
    
    chunks = []
    chunk_id = 0
    
    # Dividir por caracteres con overlap
    start = 0
    text_length = len(text)
    
    while start < text_length:
        end = start + chunk_size
        
        # Si no es el último chunk, buscar punto de corte natural (espacio o punto)
        if end < text_length:
            # Buscar el último espacio o punto dentro del chunk para cortar limpiamente
            cut_point = text.rfind(' ', start, end)
            if cut_point > start + chunk_size * 0.7:  # Solo si está en el último 30% del chunk
                end = cut_point
        
        # Extraer chunk
        chunk_text = text[start:end].strip()
        
        if chunk_text:  # Solo agregar si no está vacío
            chunks.append({
                "id": chunk_id,
                "text": chunk_text,
                "start_char": start,
                "end_char": end,
                "word_count": len(chunk_text.split()),
                "metadata": {
                    "created": datetime.now().isoformat(),
                    "chunk_size": len(chunk_text),
                    "position": f"{start}-{end}"
                }
            })
            chunk_id += 1
        
        # Avanzar con overlap
        start = end - overlap if end < text_length else text_length
    
    print(f"✅ Chunks creados: {len(chunks)}")
    return chunks


def generate_embeddings(chunks: List[Dict[str, Any]], model_name: str = EMBEDDING_MODEL) -> List[Dict[str, Any]]:
    """
    Genera embeddings para cada chunk usando modelo local.
    """
    print(f"🔢 Generando embeddings con modelo: {model_name}")
    
    # Cargar modelo de embeddings
    model = SentenceTransformer(model_name, trust_remote_code=True)
    
    # Extraer textos
    texts = [chunk["text"] for chunk in chunks]
    
    # Generar embeddings (en batches para no saturar memoria)
    embeddings = []
    batch_size = 32
    
    for i in range(0, len(texts), batch_size):
        batch = texts[i:i + batch_size]
        batch_embeddings = model.encode(batch, show_progress_bar=True)
        embeddings.extend(batch_embeddings.tolist())
        
        print(f"  Progreso: {min(i + batch_size, len(texts))}/{len(texts)} chunks procesados")
    
    # Asignar embeddings a chunks
    for i, chunk in enumerate(chunks):
        chunk["embedding"] = embeddings[i]
        chunk["embedding_dim"] = len(embeddings[i])
    
    print(f"✅ Embeddings generados: {len(chunks)} vectores de {len(embeddings[0])} dimensiones")
    return chunks


def save_to_json(chunks: List[Dict[str, Any]], output_path: str, document_info: Dict[str, Any]):
    """
    Guarda los chunks procesados en formato JSON estructurado.
    """
    print(f"💾 Guardando en: {output_path}")
    
    # Estructura final del documento
    output_data = {
        "document_info": document_info,
        "processing_info": {
            "chunk_size": CHUNK_SIZE,
            "chunk_overlap": CHUNK_OVERLAP,
            "embedding_model": EMBEDDING_MODEL,
            "processed_at": datetime.now().isoformat(),
            "total_chunks": len(chunks)
        },
        "chunks": chunks,
        "usage_instructions": {
            "how_to_query": "Usa búsqueda por similitud coseno sobre el campo 'embedding'",
            "how_to_cite": "Cita usando document_info.source y chunk.metadata",
            "recommended_top_k": 5
        }
    }
    
    # Guardar JSON
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(output_data, f, ensure_ascii=False, indent=2)
    
    # Guardar solo embeddings para FAISS (opcional, más ligero)
    embeddings_only_path = output_path.replace('.json', '_embeddings.json')
    embeddings_data = {
        "embeddings": [chunk["embedding"] for chunk in chunks],
        "chunk_ids": [chunk["id"] for chunk in chunks]
    }
    with open(embeddings_only_path, 'w', encoding='utf-8') as f:
        json.dump(embeddings_data, f, ensure_ascii=False, indent=2)
    
    file_size = os.path.getsize(output_path) / (1024 * 1024)  # MB
    print(f"✅ Archivo guardado: {output_path} ({file_size:.2f} MB)")
    print(f"✅ Embeddings guardados: {embeddings_only_path}")


def get_document_info(input_path: str) -> Dict[str, Any]:
    """
    Extrae metadatos del documento original.
    """
    return {
        "source": os.path.basename(input_path),
        "source_path": os.path.abspath(input_path),
        "source_hash": hashlib.md5(open(input_path, 'rb').read()).hexdigest(),
        "processed_filename": os.path.basename(input_path).replace('.pdf', '_processed.json'),
        "type": "pdf" if input_path.endswith('.pdf') else "text"
    }


# ================================
# FUNCIÓN PRINCIPAL
# ================================

# ================================
# FUNCIÓN PRINCIPAL ACTUALIZADA
# ================================

def process_document(input_path: str = None, output_path: str = None, book_id: str = None):
    """
    Pipeline completo de pre-procesamiento.
    
    Soporta dos modos:
    1. Modo config: book_id + carga desde YAML
    2. Modo manual: input_path + output_path explícitos
    """
    print("=" * 60)
    print("🚀 INICIANDO PRE-PROCESAMIENTO DE DOCUMENTO")
    print("=" * 60)
    
    # ================================
    # MODO CONFIG (nuevo)
    # ================================
    if book_id:
        print(f"📖 Modo config: cargando '{book_id}'")
        config = load_book_config(book_id)
        
        # Extraer valores del config
        input_path = get_config_value(config, 'source', 'file')
        output_path = get_config_value(config, 'source', 'output')
        if not output_path:
            output_path = get_config_value(config, 'source', 'file', default='').replace('.pdf', '_processed.json')
                
        # Si output_path no está en el config, construirlo
        if not output_path:
            output_dir = "data/processed"
            os.makedirs(output_dir, exist_ok=True)
            output_path = os.path.join(output_dir, f"{book_id}_processed.json")
        
        # Aplicar configuración de preprocessing
        global CHUNK_SIZE, CHUNK_OVERLAP, EMBEDDING_MODEL
        CHUNK_SIZE = get_config_value(config, 'preprocessing', 'chunk_size', default=CHUNK_SIZE)
        CHUNK_OVERLAP = get_config_value(config, 'preprocessing', 'chunk_overlap', default=CHUNK_OVERLAP)
        EMBEDDING_MODEL = get_config_value(config, 'preprocessing', 'embedding_model', default=EMBEDDING_MODEL)
        
        print(f"   - Chunk size: {CHUNK_SIZE}")
        print(f"   - Overlap: {CHUNK_OVERLAP}")
        print(f"   - Embedding model: {EMBEDDING_MODEL}")
    
    # ================================
    # MODO MANUAL (existente)
    # ================================
    elif input_path:
        print(f"📄 Modo manual: {input_path}")
        if output_path is None:
            base_name = os.path.basename(input_path).replace('.pdf', '').replace('.docx', '')
            output_dir = os.path.dirname(input_path)
            output_path = os.path.join(output_dir, f"{base_name}_processed.json")
    
    # ================================
    # VALIDACIÓN
    # ================================
    else:
        print("❌ Error: Debes especificar --input o --book")
        return False
    
    # Validar input
    if not os.path.exists(input_path):
        print(f"❌ Error: El archivo no existe: {input_path}")
        return False
    
    print(f"   - Input: {input_path}")
    print(f"   - Output: {output_path}")
    
    # ================================
    # PIPELINE (igual que antes)
    # ================================
    
    # Paso 1: Extraer texto (auto-detecta formato)
    raw_text = extract_text_from_file(input_path)
    
    # Paso 2: Limpiar texto
    clean_txt = clean_text(raw_text)
    
    # Paso 3: Crear chunks
    chunks = create_chunks(clean_txt)
    
    if not chunks:
        print("❌ Error: No se pudieron crear chunks")
        return False
    
    # Paso 4: Generar embeddings
    chunks = generate_embeddings(chunks)
    
    # Paso 5: Guardar resultado
    document_info = get_document_info(input_path)
    
    # Añadir book_id si está disponible
    if book_id:
        document_info['book_id'] = book_id
    
    save_to_json(chunks, output_path, document_info)
    
    print("=" * 60)
    print("✅ PRE-PROCESAMIENTO COMPLETADO EXITOSAMENTE")
    print("=" * 60)
    print(f"📊 Resumen:")
    print(f"   - Documento: {document_info['source']}")
    print(f"   - Chunks: {len(chunks)}")
    print(f"   - Embeddings: {chunks[0]['embedding_dim']} dimensiones")
    print(f"   - Output: {output_path}")
    print("=" * 60)
    
    return True

# ================================
# PUNTO DE ENTRADA (CLI)
# ================================

# ================================
# PUNTO DE ENTRADA (CLI) ACTUALIZADO
# ================================

if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description="Pre-procesar documentos para RAG")
    
    # Modo manual (existente)
    parser.add_argument("--input", "-i", default=None, help="Ruta al documento PDF/txt")
    parser.add_argument("--output", "-o", default=None, help="Ruta de salida JSON (opcional)")
    
    # Modo config (nuevo)
    parser.add_argument("--book", "-b", default=None, help="ID del libro (usa config YAML)")
    
    args = parser.parse_args()
    
    # Validar que se usó uno de los dos modos
    if not args.input and not args.book:
        print("❌ Error: Usa --input para modo manual o --book para modo config")
        print("Ejemplos:")
        print("  python preprocess.py --book don_quijote")
        print("  python preprocess.py --input ruta/al/documento.pdf --output ruta/output.json")
        sys.exit(1)
    
    success = process_document(
        input_path=args.input,
        output_path=args.output,
        book_id=args.book
    )
    sys.exit(0 if success else 1)